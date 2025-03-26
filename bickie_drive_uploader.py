
# === Flask Server for Saving Lesson Plans to Google Drive ===
# Required Libraries:
# pip install flask google-auth google-auth-oauthlib google-auth-httplib2 google-api-python-client python-docx

from flask import Flask, request, redirect, session, jsonify
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from docx import Document
import os
import tempfile

app = Flask(__name__)
app.secret_key = "YOUR_SECRET_KEY"
@app.route("/")
def home():
    return "Bickie uploader is running!"
@app.route("/ping")
def ping():
    return "pong"

# OAuth config
CLIENT_SECRETS_FILE = "/etc/secrets/client_secret.json"
SCOPES = ["https://www.googleapis.com/auth/drive.file"]
REDIRECT_URI = "http://localhost:5000/oauth2callback"

@app.route("/authorize")
def authorize():
    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI
    )
    auth_url, state = flow.authorization_url(access_type="offline", include_granted_scopes="true")
    session["state"] = state
    return redirect(auth_url)

@app.route("/oauth2callback")
def oauth2callback():
    state = session["state"]
    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES,
        state=state,
        redirect_uri=REDIRECT_URI
    )
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    session["credentials"] = credentials_to_dict(credentials)
    return "Google Drive connected! You can now upload."

@app.route("/save-to-drive", methods=["POST"])
def save_to_drive():
    if "credentials" not in session:
        return redirect("/authorize")

    creds = Credentials(**session["credentials"])

    lesson_title = request.json.get("lessonTitle", "LessonPlan")
    lesson_content = request.json.get("lessonContent", "")

    # Create Word Document
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = Document()
    doc.add_heading(lesson_title, 0)
    doc.add_paragraph(lesson_content)
    doc.save(temp_file.name)

    # Upload to Google Drive
    drive_service = build("drive", "v3", credentials=creds)
    file_metadata = {"name": f"{lesson_title}.docx"}
    media = MediaFileUpload(temp_file.name, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    file = drive_service.files().create(body=file_metadata, media_body=media, fields="id, webViewLink").execute()

    os.remove(temp_file.name)
    return jsonify({"message": "File uploaded!", "link": file.get("webViewLink")})

def credentials_to_dict(creds):
    return {
        "token": creds.token,
        "refresh_token": creds.refresh_token,
        "token_uri": creds.token_uri,
        "client_id": creds.client_id,
        "client_secret": creds.client_secret,
        "scopes": creds.scopes,
    }

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)
